import os
from docx import Document
import re

def remove_time_format(text):
    # Define the pattern to match the time format
    pattern = r'\b\d{2}:\d{2}\s*-\s*\d{2}:\d{2}\b'

    # Remove the time format from the text using regular expressions
    cleaned_text = re.sub(pattern, '', text)

    return cleaned_text

def copy_text_from_files(source_folder, target_folder, files_per_target):
    # Create the target folder if it doesn't exist
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)

    # Get a list of .docx files in the source folder
    files = [file for file in os.listdir(source_folder) if file.endswith('.docx')]

    # Iterate over the files and copy the text into separate target files
    for i in range(0, len(files), files_per_target):
        target_file = os.path.join(target_folder, f"target_{i // files_per_target + 1}.docx")

        # Create a new document for each target file
        target_doc = Document()

        # Iterate over the files for each target and copy the text
        for file in files[i:i+files_per_target]:
            file_path = os.path.join(source_folder, file)
            source_doc = Document(file_path)

            # Iterate over paragraphs in the source document
            for paragraph in source_doc.paragraphs:
                # Remove time format from the paragraph text
                cleaned_text = remove_time_format(paragraph.text)

                # Copy the paragraph to the target document if it's not empty
                if cleaned_text.strip():
                    target_doc.add_paragraph(cleaned_text, style=paragraph.style)

        # Save the target document
        target_doc.save(target_file)

    print("Text copied successfully!")

# Example usage
source_folder = r"C:\test\source"
target_folder = r"C:\test\target"
files_per_target = 50

copy_text_from_files(source_folder, target_folder, files_per_target)
