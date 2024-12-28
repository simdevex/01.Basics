import os
from docx import Document

def copy_text(source_doc, target_doc):
    # Open source document
    source = Document(source_doc)

    # Iterate through paragraphs in the source document
    for paragraph in source.paragraphs:
        # Copy the paragraph with formatting to the target document
        target_doc.add_paragraph(paragraph.text, style=paragraph.style)

def merge_files(source_folder, target_folder):
    # Create the target folder if it doesn't exist
    os.makedirs(target_folder, exist_ok=True)

    file_count = 0
    target_count = 1
    target_doc = Document()

    # Iterate through the .docx files in the source folder
    for filename in os.listdir(source_folder):
        if filename.endswith(".docx"):
            source_file = os.path.join(source_folder, filename)
            copy_text(source_file, target_doc)
            file_count += 1

            # Create a new target file after every 50 source files
            if file_count % 50 == 0:
                target_file = os.path.join(target_folder, f"target{target_count}.docx")
                target_doc.save(target_file)
                print(f"Merging of files {file_count-49}-{file_count} complete. The target file is saved at:", target_file)
                target_count += 1
                target_doc = Document()

    # Save the remaining text in the last target file
    if len(target_doc.paragraphs) > 0:
        target_file = os.path.join(target_folder, f"target{target_count}.docx")
        target_doc.save(target_file)
        print(f"Merging of files {file_count-(file_count%50)+1}-{file_count} complete. The target file is saved at:", target_file)

if __name__ == "__main__":
    source_folder = "C:\\test\\source"  # Specify the folder path containing the source .docx files
    target_folder = "C:\\test\\target"  # Specify the folder path where the target .docx files will be created

    merge_files(source_folder, target_folder)
